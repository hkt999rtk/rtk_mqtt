#!/usr/bin/env python3
"""
MQTT Diagnostic Communication Protocol Specification Document Generator
Converts SPEC.md into a formatted Word document with automatic TOC
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.oxml.shared import qn
import markdown
from bs4 import BeautifulSoup

def setup_document_styles(doc):
    """Setup custom styles for the document"""
    styles = doc.styles
    
    # Title style
    if 'Custom Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(20)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
    
    # Modify built-in heading styles for TOC compatibility
    heading_sizes = [16, 14, 12, 11, 10]
    for i, size in enumerate(heading_sizes, 1):
        style_name = f'Heading {i}'
        if style_name in [s.name for s in styles]:
            heading_style = styles[style_name]
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.size = Pt(size)
            heading_font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    # Code block style
    if 'Code Block' not in [s.name for s in styles]:
        code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(8)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    
    # Custom blockquote style
    if 'Custom Quote' not in [s.name for s in styles]:
        quote_style = styles.add_style('Custom Quote', WD_STYLE_TYPE.PARAGRAPH)
        quote_font = quote_style.font
        quote_font.name = 'Times New Roman'
        quote_font.size = Pt(10)
        quote_font.italic = True
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.space_before = Pt(6)
        quote_style.paragraph_format.space_after = Pt(6)

def add_header_and_footer(doc):
    """Add header with logo and title, and footer with page numbers"""
    from pathlib import Path
    
    # Get the last section (not the title page)
    section = doc.sections[-1]
    
    # === Add Header ===
    header = section.header
    header_para = header.paragraphs[0]
    
    # Create a table in header for layout (logo left, title right)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(1.5)  # Logo column
    header_table.columns[1].width = Inches(5.0)  # Title column
    
    # Left cell - Logo
    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    
    # Try to add logo
    logo_path = Path(__file__).parent / "page-logo.png"
    if logo_path.exists():
        try:
            logo_run = left_para.add_run()
            logo_run.add_picture(str(logo_path), width=Inches(1.0))
        except Exception as e:
            # If logo fails, add placeholder text
            left_para.add_run("[Logo]").font.size = Pt(8)
    else:
        left_para.add_run("[Logo]").font.size = Pt(8)
    
    # Right cell - Title
    right_cell = header_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = right_para.add_run("MQTT Diagnostic 通訊規格")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(10)
    title_run.font.bold = True
    
    # Remove table borders
    for row in header_table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            # Set all borders to none
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
    
    # === Add Footer ===
    footer = section.footer
    
    # Create a paragraph in the footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create the page number field
    run = footer_para.add_run()
    
    # Add the PAGE field
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    # Format the page number
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def add_table_of_contents(doc):
    """Add a table of contents placeholder"""
    # Add TOC title
    toc_title = doc.add_paragraph()
    toc_title.style = 'Heading 1'
    toc_title_run = toc_title.runs[0] if toc_title.runs else toc_title.add_run()
    toc_title_run.text = "目錄"
    
    # Add TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # Insert TOC field code
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    # Add page break
    doc.add_page_break()

def parse_markdown_content(content):
    """Parse markdown content and return structured data"""
    lines = content.split('\n')
    sections = []
    current_section = {'level': 0, 'title': '', 'content': []}
    
    for line in lines:
        # Check for headers
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            # Save previous section if it has content
            if current_section['title'] or current_section['content']:
                sections.append(current_section)
            
            # Start new section
            level = len(header_match.group(1))
            title = header_match.group(2).strip()
            current_section = {'level': level, 'title': title, 'content': []}
        else:
            current_section['content'].append(line)
    
    # Add the last section
    if current_section['title'] or current_section['content']:
        sections.append(current_section)
    
    return sections

def add_markdown_table_to_doc(doc, table_lines):
    """Convert markdown table to Word table"""
    if not table_lines:
        return
    
    # Parse table
    rows = []
    for line in table_lines:
        if '|' in line:
            # Check if this is a separator line (contains only |, -, :, and whitespace)
            separator_pattern = r'^\s*\|?[\s\-\|:]+\|?\s*$'
            if re.match(separator_pattern, line):
                continue  # Skip separator lines
            
            # Split by | and clean up
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty first/last cells if they exist
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            if cells:
                rows.append(cells)
    
    if not rows:
        return
    
    # Create Word table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Set table alignment
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fill table
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < len(table.rows[i].cells):
                # Clean up cell content
                cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_data)  # Remove bold
                cell_text = re.sub(r'\*(.*?)\*', r'\1', cell_text)      # Remove italic
                cell_text = re.sub(r'`(.*?)`', r'\1', cell_text)        # Remove code
                
                cell = table.rows[i].cells[j]
                cell.text = cell_text.strip()
                
                # Format cell
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Set font size for table cells
                if cell_para.runs:
                    for run in cell_para.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        # Make header row bold
                        if i == 0:
                            run.font.bold = True
                else:
                    # If no runs, create one with the text
                    run = cell_para.add_run(cell_text.strip())
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    if i == 0:
                        run.font.bold = True
                    cell.text = ""  # Clear the text since we added it to run

def add_code_block_to_doc(doc, code_lines, language=''):
    """Add code block to document"""
    if not code_lines:
        return
    
    # Join code lines
    code_text = '\n'.join(code_lines)
    
    # Add code block
    code_para = doc.add_paragraph()
    code_para.style = 'Code Block'
    
    # Add the code text
    run = code_para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

def add_content_to_doc(doc, sections, base_path=None):
    """Add parsed content to Word document"""
    if base_path is None:
        base_path = Path(__file__).parent
    
    for section in sections:
        # 有標題才加標題
        if section['title']:
            if section['level'] <= 5:
                heading_para = doc.add_paragraph()
                heading_para.style = f'Heading {min(section["level"], 5)}'
                heading_run = heading_para.add_run(section['title'])
        
        # 處理內容
        content_lines = section['content']
        current_para = None
        table_lines = []
        code_lines = []
        in_table = False
        in_code_block = False
        
        i = 0
        while i < len(content_lines):
            line = content_lines[i]
            
            # Check for code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block
                    add_code_block_to_doc(doc, code_lines)
                    code_lines = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                    code_lines = []
                i += 1
                continue
            
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            line = line.strip()
            
            # Check if line is part of a table
            if '|' in line and line:
                in_table = True
                table_lines.append(line)
            elif in_table and line == '':
                # End of table
                add_markdown_table_to_doc(doc, table_lines)
                table_lines = []
                in_table = False
            elif in_table:
                # Continue table or end it
                if '|' in line:
                    table_lines.append(line)
                else:
                    # End table and process this line
                    add_markdown_table_to_doc(doc, table_lines)
                    table_lines = []
                    in_table = False
                    # Process current line as regular content
                    if line:
                        para = doc.add_paragraph()
                        run = para.add_run(line)
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
            else:
                # Regular content
                if line:
                    # Check for image references
                    img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
                    if img_match:
                        alt_text = img_match.group(1)
                        img_path = img_match.group(2)
                        
                        # Create paragraph for image
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Try to find and add the image
                        full_img_path = base_path / img_path
                        if full_img_path.exists():
                            try:
                                img_para.add_run().add_picture(str(full_img_path), width=Inches(3.0))
                            except Exception as e:
                                # If image fails to load, add alt text
                                img_para.add_run(f"[圖片: {alt_text}]").font.italic = True
                        else:
                            # Image not found, add placeholder
                            img_para.add_run(f"[圖片: {alt_text}]").font.italic = True
                        
                        # Skip the rest of the processing for image lines
                        i += 1
                        continue
                    
                    # Check for horizontal rules (--- or more)
                    if re.match(r'^-{3,}$', line.strip()):
                        # Just skip horizontal rules as they are section separators
                        i += 1
                        continue
                    
                    # Skip empty or whitespace-only lines
                    if not line.strip():
                        i += 1
                        continue
                    
                    # Regular paragraph handling
                    para = doc.add_paragraph()
                    
                    # Handle different formatting
                    if line.startswith('- ') or line.startswith('* '):
                        para.style = 'List Bullet'
                        text = line[2:].strip()
                    elif re.match(r'^\d+\.', line):
                        para.style = 'List Number'
                        # Extract text after the number
                        text = re.sub(r'^\d+\.\s*', '', line)
                    elif line.startswith('>'):
                        # Blockquote
                        para.style = 'Custom Quote'
                        text = line[1:].strip()
                    else:
                        text = line
                    
                    # Handle inline formatting
                    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
                    
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            # Bold text
                            run = para.add_run(part[2:-2])
                            run.font.bold = True
                        elif part.startswith('*') and part.endswith('*'):
                            # Italic text
                            run = para.add_run(part[1:-1])
                            run.font.italic = True
                        elif part.startswith('`') and part.endswith('`'):
                            # Inline code
                            run = para.add_run(part[1:-1])
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        else:
                            # Regular text
                            run = para.add_run(part)
                        
                        # Set default font properties
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
            
            i += 1
        
        # Handle any remaining table or code block
        if table_lines:
            add_markdown_table_to_doc(doc, table_lines)
        if code_lines:
            add_code_block_to_doc(doc, code_lines)

def create_mqtt_spec_document():
    """Main function to create the MQTT specification document"""
    print("🚀 開始產生 MQTT Diagnostic 通訊規格文件...")
    
    # Create new document
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc)
    
    # Add title page with logo
    # Add some vertical space before logo
    doc.add_paragraph().space_before = Pt(72)  # 1 inch spacing
    
    # Add Realtek logo to title page
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cover_logo_path = Path(__file__).parent / "cover-logo.png"
    if cover_logo_path.exists():
        try:
            logo_para.add_run().add_picture(str(cover_logo_path), width=Inches(4.0))
        except Exception as e:
            logo_para.add_run("[Realtek Logo]").font.size = Pt(16)
    else:
        logo_para.add_run("[Realtek Logo]").font.size = Pt(16)
    
    # Add some space after logo
    doc.add_paragraph().space_before = Pt(36)
    
    # Add title
    title_para = doc.add_paragraph()
    title_para.style = 'Custom Title'
    title_run = title_para.add_run("MQTT Diagnostic 通訊規格")
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.add_run("設備診斷與監控通訊協議規格書").font.size = Pt(12)
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.space_before = Pt(24)
    info_text = "適用範圍: WiFi 診斷事件、設備健康監控\n通訊協議: MQTT v3.1.1 / v5.0\n版本: v1.0"
    info_para.add_run(info_text).font.size = Pt(10)
    
    # Create a new section for pages after title page (with headers/footers)
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Set the first section (title page) to have different first page header
    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = True
    
    # Add header and footer to the new section (applies to all pages except title page)
    add_header_and_footer(doc)
    
    # Add table of contents
    add_table_of_contents(doc)
    
    # Process main SPEC.md file first
    spec_path = Path(__file__).parent.parent / "SPEC.md"
    if not spec_path.exists():
        print(f"❌ 找不到 SPEC.md 檔案: {spec_path}")
        return None
    
    print(f"📝 處理主要規格檔案: {spec_path.name}")
    
    try:
        with open(spec_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse markdown content
        sections = parse_markdown_content(content)
        
        # Add content to document
        add_content_to_doc(doc, sections, spec_path.parent)
        
    except Exception as e:
        print(f"❌ 處理 {spec_path.name} 時發生錯誤: {e}")
        return None
    
    # Add page break before appendices
    doc.add_page_break()
    
    # Add appendices heading
    appendix_title = doc.add_paragraph()
    appendix_title.style = 'Heading 1'
    appendix_title_run = appendix_title.add_run("附錄")
    
    # Process additional markdown files as appendices
    # 順序安排：理論概念 -> 通訊流程 -> 實作範例 (由簡到複雜，常用到不常用)
    additional_files = [
        # 基礎理解：系統架構和通訊流程
        ("architecture_diagram.txt", "A. 系統架構與元件關係"),
        ("mqtt_flow_diagrams.md", "B. MQTT 通訊流程與時序圖"),
        
        # 實作參考：從常見到複雜的診斷場景
        ("connect_failure_example.md", "C. WiFi 連線失敗診斷範例"),  # 最常見的問題
        ("arp_loss_example.md", "D. ARP 遺失診斷範例"),           # 網路連線問題 
        ("roaming_diagnosis_example.md", "E. WiFi 漫遊診斷範例")    # 較複雜的診斷場景
    ]
    
    base_path = Path(__file__).parent.parent
    
    for filename, appendix_title in additional_files:
        file_path = base_path / filename
        if file_path.exists():
            print(f"📝 處理附錄檔案: {filename}")
            
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Add appendix title with description
                appendix_heading = doc.add_paragraph()
                appendix_heading.style = 'Heading 2'
                appendix_heading_run = appendix_heading.add_run(appendix_title)
                
                # Add brief description for each appendix
                description_map = {
                    "A. 系統架構與元件關係": "系統整體架構圖，說明 Device、Controller、MQTT Broker 之間的關係與資料流向",
                    "B. MQTT 通訊流程與時序圖": "詳細的 MQTT 訊息交換時序圖，包含各種診斷場景的完整通訊流程",
                    "C. WiFi 連線失敗診斷範例": "最常見的連線問題診斷，包含認證失敗、超時等處理流程",
                    "D. ARP 遺失診斷範例": "網路連線品質診斷，包含封包遺失檢測與網路環境分析",
                    "E. WiFi 漫遊診斷範例": "複雜的漫遊機制診斷，包含 AP 選擇與切換過程分析"
                }
                
                if appendix_title in description_map:
                    desc_para = doc.add_paragraph()
                    desc_run = desc_para.add_run(description_map[appendix_title])
                    desc_run.font.size = Pt(9)
                    desc_run.font.name = 'Times New Roman'
                    desc_run.font.italic = True
                    desc_para.space_after = Pt(12)
                
                # Parse and add content
                if filename.endswith('.txt'):
                    # Handle plain text files (like architecture diagram)
                    lines = content.split('\n')
                    for line in lines:
                        if line.strip():
                            para = doc.add_paragraph()
                            run = para.add_run(line)
                            run.font.size = Pt(9)
                            run.font.name = 'Courier New'
                else:
                    # Handle markdown files
                    sections = parse_markdown_content(content)
                    add_content_to_doc(doc, sections, base_path)
                
                # Add page break between appendices (except for last one)
                if filename != additional_files[-1][0]:
                    doc.add_page_break()
                    
            except Exception as e:
                print(f"❌ 處理附錄 {filename} 時發生錯誤: {e}")
                continue
        else:
            print(f"⚠️  找不到附錄檔案: {filename}")
            continue
    
    # Save document
    output_path = Path(__file__).parent / "MQTT_Diagnosis_通訊規格.docx"
    doc.save(output_path)
    
    print(f"✅ 文件已儲存至: {output_path}")
    print("📋 注意: 要更新目錄，請在 Microsoft Word 中:")
    print("   1. 開啟文件")
    print("   2. 在目錄上按右鍵")
    print("   3. 選擇 '更新功能變數' -> '更新整個目錄'")
    
    return output_path

if __name__ == "__main__":
    try:
        output_file = create_mqtt_spec_document()
        if output_file:
            print(f"\n🎉 成功產生 MQTT 規格文件!")
            print(f"📁 輸出檔案: {output_file}")
    except Exception as e:
        print(f"💥 錯誤: {e}")
        import traceback
        traceback.print_exc()