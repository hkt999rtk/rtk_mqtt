#!/usr/bin/env python3
"""
HA IoT MCP Document Generator
Converts HA_IOT_MCP.md into a formatted Word document with images
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.oxml.shared import qn

def setup_document_styles(doc):
    """Setup custom styles for the document"""
    styles = doc.styles
    
    # Title style
    if 'Custom Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(18)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
    
    # Modify built-in heading styles
    heading_sizes = [16, 14, 12, 11, 10]
    for i, size in enumerate(heading_sizes, 1):
        style_name = f'Heading {i}'
        if style_name in [s.name for s in styles]:
            heading_style = styles[style_name]
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.size = Pt(size)
            heading_font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    # Code block style
    if 'Code Block' not in [s.name for s in styles]:
        code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    
    # Custom blockquote style
    if 'Custom Quote' not in [s.name for s in styles]:
        quote_style = styles.add_style('Custom Quote', WD_STYLE_TYPE.PARAGRAPH)
        quote_font = quote_style.font
        quote_font.name = 'Times New Roman'
        quote_font.size = Pt(10)
        quote_font.italic = True
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.space_before = Pt(6)
        quote_style.paragraph_format.space_after = Pt(6)

def add_header_and_footer(doc):
    """Add header with logo and title, and footer with page numbers"""
    from pathlib import Path
    
    # Get the last section (not the title page)
    section = doc.sections[-1]
    
    # === Add Header ===
    header = section.header
    header_para = header.paragraphs[0]
    
    # Create a table in header for layout (logo left, title right)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(1.5)  # Logo column
    header_table.columns[1].width = Inches(5.0)  # Title column
    
    # Left cell - Logo
    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    
    # Try to add logo
    logo_path = Path(__file__).parent / "page-logo.png"
    if logo_path.exists():
        try:
            logo_run = left_para.add_run()
            logo_run.add_picture(str(logo_path), width=Inches(1.0))
        except Exception as e:
            # If logo fails, add placeholder text
            left_para.add_run("[Logo]").font.size = Pt(8)
    else:
        left_para.add_run("[Logo]").font.size = Pt(8)
    
    # Right cell - Title
    right_cell = header_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = right_para.add_run("HA IoT MCP 整合架構")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(10)
    title_run.font.bold = True
    
    # Remove table borders
    for row in header_table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            # Set all borders to none
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
    
    # Add Footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number field
    run = footer_para.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def add_table_of_contents(doc):
    """Add a table of contents"""
    toc_title = doc.add_paragraph()
    toc_title.style = 'Heading 1'
    toc_title_run = toc_title.runs[0] if toc_title.runs else toc_title.add_run()
    toc_title_run.text = "目錄"
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    doc.add_page_break()

def parse_markdown_content(content):
    """Parse markdown content and return structured data"""
    lines = content.split('\n')
    sections = []
    current_section = {'level': 0, 'title': '', 'content': []}
    
    for line in lines:
        # Check for headers
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            # Save previous section if it has content
            if current_section['title'] or current_section['content']:
                sections.append(current_section)
            
            # Start new section
            level = len(header_match.group(1))
            title = header_match.group(2).strip()
            current_section = {'level': level, 'title': title, 'content': []}
        else:
            current_section['content'].append(line)
    
    # Add the last section
    if current_section['title'] or current_section['content']:
        sections.append(current_section)
    
    return sections

def add_code_block_to_doc(doc, code_lines):
    """Add code block to document"""
    if not code_lines:
        return
    
    code_text = '\n'.join(code_lines)
    code_para = doc.add_paragraph()
    code_para.style = 'Code Block'
    
    run = code_para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def add_content_to_doc(doc, sections, base_path=None):
    """Add parsed content to Word document"""
    if base_path is None:
        base_path = Path(__file__).parent.parent
    
    for section in sections:
        # Add heading if exists
        if section['title']:
            if section['level'] <= 5:
                heading_para = doc.add_paragraph()
                heading_para.style = f'Heading {min(section["level"], 5)}'
                heading_run = heading_para.add_run(section['title'])
        
        # Process content
        content_lines = section['content']
        code_lines = []
        in_code_block = False
        
        i = 0
        while i < len(content_lines):
            line = content_lines[i]
            
            # Check for code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block
                    add_code_block_to_doc(doc, code_lines)
                    code_lines = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                    code_lines = []
                i += 1
                continue
            
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            line = line.strip()
            
            if line:
                # Check for image references
                img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
                if img_match:
                    alt_text = img_match.group(1)
                    img_path = img_match.group(2)
                    
                    # Create paragraph for image
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Try to find and add the image
                    full_img_path = base_path / img_path
                    if full_img_path.exists():
                        try:
                            img_para.add_run().add_picture(str(full_img_path), width=Inches(5.0))
                            # Add caption below image
                            caption_para = doc.add_paragraph()
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_run = caption_para.add_run(f"圖: {alt_text}")
                            caption_run.font.size = Pt(9)
                            caption_run.font.name = 'Times New Roman'
                            caption_run.font.italic = True
                        except Exception as e:
                            img_para.add_run(f"[圖片載入失敗: {alt_text}]").font.italic = True
                    else:
                        img_para.add_run(f"[圖片未找到: {alt_text}]").font.italic = True
                    
                    i += 1
                    continue
                
                # Skip horizontal rules
                if re.match(r'^-{3,}$', line.strip()):
                    i += 1
                    continue
                
                # Regular paragraph handling
                para = doc.add_paragraph()
                
                # Handle different formatting
                if line.startswith('- ') or line.startswith('* '):
                    para.style = 'List Bullet'
                    text = line[2:].strip()
                elif re.match(r'^\d+\.', line):
                    para.style = 'List Number'
                    text = re.sub(r'^\d+\.\s*', '', line)
                elif line.startswith('>'):
                    para.style = 'Custom Quote'
                    text = line[1:].strip()
                else:
                    text = line
                
                # Handle inline formatting
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
                
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        run = para.add_run(part[2:-2])
                        run.font.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        # Italic text
                        run = para.add_run(part[1:-1])
                        run.font.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Inline code
                        run = para.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    else:
                        # Regular text
                        run = para.add_run(part)
                    
                    # Set default font properties
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
            
            i += 1
        
        # Handle any remaining code block
        if code_lines:
            add_code_block_to_doc(doc, code_lines)

def create_ha_iot_document():
    """Main function to create the HA IoT MCP document"""
    print("🚀 開始產生 HA IoT MCP 整合架構文件...")
    
    # Create new document
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc)
    
    # Add title page with logo
    # Add some vertical space before logo
    doc.add_paragraph().space_before = Pt(72)  # 1 inch spacing
    
    # Add Realtek logo to title page
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cover_logo_path = Path(__file__).parent / "cover-logo.png"
    if cover_logo_path.exists():
        try:
            logo_para.add_run().add_picture(str(cover_logo_path), width=Inches(4.0))
        except Exception as e:
            logo_para.add_run("[Realtek Logo]").font.size = Pt(16)
    else:
        logo_para.add_run("[Realtek Logo]").font.size = Pt(16)
    
    # Add some space after logo
    doc.add_paragraph().space_before = Pt(36)
    
    # Add title
    title_para = doc.add_paragraph()
    title_para.style = 'Custom Title'
    title_run = title_para.add_run("Home Assistant / MQTT / IoT Device")
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.add_run("MCP / Claude Desktop 整合架構").font.size = Pt(14)
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.space_before = Pt(24)
    info_text = "智能家居控制系統整合文件\n適用於 AI 助手與物聯網設備控制\n版本: v1.0"
    info_para.add_run(info_text).font.size = Pt(11)
    
    # Create new section for content pages
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Set title page to have different header/footer
    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = True
    
    # Add header and footer
    add_header_and_footer(doc)
    
    # Add table of contents
    add_table_of_contents(doc)
    
    # Process HA_IOT_MCP.md file
    md_path = Path(__file__).parent.parent / "HA_IOT_MCP.md"
    if not md_path.exists():
        print(f"❌ 找不到 HA_IOT_MCP.md 檔案: {md_path}")
        return None
    
    print(f"📝 處理文件: {md_path.name}")
    
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse markdown content
        sections = parse_markdown_content(content)
        
        # Add content to document
        add_content_to_doc(doc, sections, md_path.parent)
        
    except Exception as e:
        print(f"❌ 處理 {md_path.name} 時發生錯誤: {e}")
        return None
    
    # Save document
    output_path = Path(__file__).parent / "HA_IoT_MCP_整合架構.docx"
    doc.save(output_path)
    
    print(f"✅ 文件已儲存至: {output_path}")
    print("📋 注意: 要更新目錄，請在 Microsoft Word 中:")
    print("   1. 開啟文件")
    print("   2. 在目錄上按右鍵")
    print("   3. 選擇 '更新功能變數' -> '更新整個目錄'")
    
    return output_path

if __name__ == "__main__":
    try:
        output_file = create_ha_iot_document()
        if output_file:
            print(f"\n🎉 成功產生 HA IoT MCP 整合架構文件!")
            print(f"📁 輸出檔案: {output_file}")
    except Exception as e:
        print(f"💥 錯誤: {e}")
        import traceback
        traceback.print_exc()